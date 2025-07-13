import streamlit as st
from docx import Document
from fpdf import FPDF
import io
import pandas as pd


def generate_docx(summary_text, table_df=None):
    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading("Sample Size Report", 0)
    
    for line in summary_text.split('\n'):
        doc.add_paragraph(line)
    
    if table_df is not None:
        doc.add_paragraph("Sample Size Table")
        table = doc.add_table(rows=1, cols=len(table_df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(table_df.columns):
            hdr_cells[i].text = str(col)
        for _, row in table_df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf(summary_text, table_df=None):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Add summary lines
    for line in summary_text.split('\n'):
        pdf.cell(200, 10, txt=line, ln=True)

    # Add table if present
    if table_df is not None:
        pdf.ln(10)
        pdf.cell(200, 10, txt="Sample Size Table", ln=True)
        col_names = [str(col) for col in table_df.columns]
        pdf.set_font("Arial", 'B', size=10)
        pdf.cell(0, 10, ' | '.join(col_names), ln=True)
        pdf.set_font("Arial", size=10)
        for _, row in table_df.iterrows():
            row_str = ' | '.join(str(item) for item in row)
            pdf.cell(0, 10, row_str, ln=True)

    # Convert to binary and return BytesIO object
    buffer = io.BytesIO()
    pdf_output = pdf.output(dest='S').encode('latin1')  # <-- capture PDF as string
    buffer.write(pdf_output)
    buffer.seek(0)
    return buffer

def download_report(summary_text, table_df=None, filename="sample_size_report", filetype="docx"):
    """
    Universal Streamlit download button for PDF or DOCX reports.
    """
    if filetype == "docx":
        buffer = generate_docx(summary_text, table_df)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif filetype == "pdf":
        buffer = generate_pdf(summary_text, table_df)
        mime_type = "application/pdf"
        ext = "pdf"
    else:
        st.error("Unsupported file type.")
        return

    st.download_button(
        label=f"📄 Download Report as {ext.upper()}",
        data=buffer,
        file_name=f"{filename}.{ext}",
        mime=mime_type
    )
